"""Word .docx 文档解析器。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from day04_app.common.exceptions import BusinessException
from day04_app.schemas.knowledge_schema import ParsedDocument, ParsedDocumentSegment


class WordDocumentParser:
    """将 Word 段落和表格行转换为可追溯的统一文本段。"""

    supported_file_types = frozenset({"docx"})

    def parse(self, *, document_id: str, file_path: Path) -> ParsedDocument:
        try:
            word_document = Document(file_path)
        except (PackageNotFoundError, ValueError, OSError) as exc:
            raise BusinessException(code=40058, message="Word 文档无法读取或文件已损坏") from exc

        segments: list[ParsedDocumentSegment] = []
        for paragraph_index, paragraph in enumerate(word_document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            segments.append(
                ParsedDocumentSegment(
                    segment_index=len(segments),
                    text=text,
                    location=f"Paragraph:{paragraph_index}",
                    metadata={"source_type": "paragraph", "paragraph_index": paragraph_index},
                )
            )

        for table_index, table in enumerate(word_document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                # 单元格按列拼接；空单元格不进入知识库，防止生成无意义的空白段。
                cell_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cell_values:
                    continue
                segments.append(
                    ParsedDocumentSegment(
                        segment_index=len(segments),
                        text=" | ".join(cell_values),
                        location=f"Table:{table_index}/Row:{row_index}",
                        metadata={
                            "source_type": "table_row",
                            "table_index": table_index,
                            "row_index": row_index,
                        },
                    )
                )

        if not segments:
            raise BusinessException(code=40059, message="Word 文档中没有可用于知识库的文本或表格内容")

        return ParsedDocument(
            document_id=document_id,
            # 上传元数据表尚未接入时，物理文件名仅用于本阶段调试，不会暴露给前端。
            file_name=file_path.name,
            file_type="docx",
            parser_name=self.__class__.__name__,
            segments=segments,
            metadata={
                "paragraph_count": len(word_document.paragraphs),
                "table_count": len(word_document.tables),
                "segment_count": len(segments),
            },
        )
